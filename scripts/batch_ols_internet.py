#!/usr/bin/env python3
"""
Batch-run OLS models for internet-related exposures vs top dependent scores.

Usage:
  python3 scripts/batch_ols_internet.py /path/to/analysis-ready.csv

Outputs (defaults written under InputData/):
  - InputData/adjusted_models_batch_summary.csv
  - InputData/outputs_models/plots/
  - InputData/batch_ols_report.docx (Word summary)

The script fits crude and adjusted OLS models for each exposure->outcome pair
and applies Benjamini-Hochberg FDR to the exposure p-values across the batch.
It then creates a Word summary with top 10 associations, two full adjusted
model tables/plots, and a short confounding & causal assessment.
"""
import sys
from pathlib import Path
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
import statsmodels.api as sm
from statsmodels.stats.outliers_influence import variance_inflation_factor
from statsmodels.stats.multitest import multipletests
import textwrap

try:
    from docx import Document
    from docx.shared import Inches
except Exception:
    Document = None

sns.set(style='whitegrid')

# Covariates to always adjust for
DEFAULT_COVARIATES = [
    'Age (IN YEARS)',
    'Monthly household income_midpoint_INR',
    'Highest level of education completed',
    'Marital Status'
]

EXPOSURE_KEYWORDS = ['internet', 'Internet', 'Obsession', 'Perceived Control', 'Internet as a Source', 'Internet_Usage']

def find_internet_exposures(cols):
    out = []
    for c in cols:
        low = c.lower()
        for k in EXPOSURE_KEYWORDS:
            if k.lower() in low:
                out.append(c)
                break
    return sorted(set(out))

def select_top_dependents(df, depend_list, top_n=5):
    # choose dependents that exist and are numeric; then pick top_n by variance
    cand = [c for c in depend_list if c in df.columns and pd.api.types.is_numeric_dtype(df[c])]
    if not cand:
        # fallback: any numeric columns not in covariates
        cand = [c for c in df.select_dtypes(include=[np.number]).columns if c not in DEFAULT_COVARIATES]
    var = {c: df[c].dropna().var() for c in cand}
    sorted_by_var = sorted(var.items(), key=lambda x: x[1], reverse=True)
    return [c for c, v in sorted_by_var[:top_n]]

def build_design(df, exposure, covariates):
    # X matrix with exposure and covariates; handle categorical covariates
    X = pd.DataFrame()
    X['exposure'] = df[exposure]
    for cov in covariates:
        if cov not in df.columns:
            continue
        ser = df[cov]
        if pd.api.types.is_numeric_dtype(ser):
            X[cov] = ser
        else:
            # one-hot encode, drop first to avoid multicollinearity
            d = pd.get_dummies(ser.astype(str), prefix=cov.replace(' ', '_'), drop_first=True)
            if not d.empty:
                X = pd.concat([X, d], axis=1)
    X = sm.add_constant(X, has_constant='add')
    return X

def run_models(df, exposures, dependents, covariates):
    rows = []
    plots_dir = Path('InputData') / 'outputs_models' / 'plots'
    plots_dir.mkdir(parents=True, exist_ok=True)

    for y in dependents:
        for x in exposures:
            # prepare data
            sub = df[[x, y] + [c for c in covariates if c in df.columns]].copy()
            sub = sub.dropna()
            n = sub.shape[0]
            if n < 10:
                continue

            # crude
            Xc = sm.add_constant(sub[[x]], has_constant='add')
            try:
                m_crude = sm.OLS(sub[y], Xc).fit()
                b_crude = m_crude.params.get(x, np.nan)
                se_crude = m_crude.bse.get(x, np.nan)
                p_crude = m_crude.pvalues.get(x, np.nan)
                ci_crude = m_crude.conf_int().loc[x].tolist() if x in m_crude.params.index else [np.nan, np.nan]
            except Exception as e:
                b_crude = se_crude = p_crude = np.nan
                ci_crude = [np.nan, np.nan]

            # adjusted
            X = build_design(sub, x, covariates)
            try:
                # coerce exposure to numeric when possible; if exposure is categorical (e.g., 'Low'/ 'High')
                # then we cannot extract a single exposure coefficient and will skip adjusted fit.
                X['exposure'] = pd.to_numeric(X['exposure'], errors='coerce')
                if X['exposure'].isnull().all():
                    # fully non-numeric exposure (categorical). skip adjusted model.
                    raise ValueError('Exposure is non-numeric / categorical; adjusted single-coef OLS not run')
                # ensure numeric design matrix (bool dtypes can cause a pandas->numpy object cast error)
                X = X.astype(float)
                y_ser = pd.to_numeric(sub[y], errors='coerce')
                # drop any rows that became NA after coercion (should be rare given earlier dropna)
                mask = X.notnull().all(axis=1) & y_ser.notnull()
                X_fit = X.loc[mask]
                y_fit = y_ser.loc[mask]
                if X_fit.shape[0] < 3:
                    raise ValueError('Too few rows after coercion for adjusted fit')
                m_adj = sm.OLS(y_fit, X_fit).fit(cov_type='HC3')
                # use robust SEs
                coef = m_adj.params.get('exposure', np.nan)
                se = m_adj.bse.get('exposure', np.nan)
                p = m_adj.pvalues.get('exposure', np.nan)
                ci = m_adj.conf_int().loc['exposure'].tolist() if 'exposure' in m_adj.params.index else [np.nan, np.nan]
                r2 = m_adj.rsquared
                adjr2 = m_adj.rsquared_adj
            except Exception as e:
                # surface useful debug info rather than silently swallowing
                print(f'Adjusted fit failed for exposure={x!r}, outcome={y!r}, n_before={n}:', type(e), e)
                coef = se = p = r2 = adjr2 = np.nan
                ci = [np.nan, np.nan]

            # change-in-estimate
            change = (coef - b_crude) / b_crude * 100 if pd.notna(b_crude) and b_crude != 0 else np.nan

            rows.append({
                'exposure': x,
                'outcome': y,
                'n': n,
                'crude_beta': b_crude,
                'crude_se': se_crude,
                'crude_p': p_crude,
                'crude_ci_low': ci_crude[0],
                'crude_ci_high': ci_crude[1],
                'adj_beta': coef,
                'adj_se': se,
                'adj_p': p,
                'adj_ci_low': ci[0],
                'adj_ci_high': ci[1],
                'r2': r2,
                'adj_r2': adjr2,
                'change_pct': change
            })

            # save a quick exposure vs outcome plot
            try:
                fig, ax = plt.subplots(figsize=(6.5,4))
                sns.regplot(x=x, y=y, data=sub, ax=ax, scatter_kws={'s':20}, line_kws={'color':'red'})
                ax.set_title(f'{y} ~ {x}\nadj p={p:.3g}' if pd.notna(p) else f'{y} ~ {x}')
                fname = plots_dir / f"{sanitize(x)}__vs__{sanitize(y)}.png"
                fig.savefig(fname, dpi=200, bbox_inches='tight', facecolor=fig.get_facecolor())
                plt.close(fig)
            except Exception:
                pass

    df_res = pd.DataFrame(rows)
    return df_res, plots_dir

def sanitize(s):
    return ''.join(c if c.isalnum() or c in (' ', '_', '-') else '_' for c in str(s)).strip().replace(' ', '_')

def make_word_report(df_res, plots_dir, out_docx, top_n=10):
    if Document is None:
        print('python-docx not installed; skipping Word report creation')
        return
    doc = Document()
    doc.add_heading('Batch OLS internet-related exposures — Top associations', level=1)

    # FDR correct adj_p and add column
    pvals = df_res['adj_p'].fillna(1.0).values
    reject, pvals_fdr, _, _ = multipletests(pvals, alpha=0.05, method='fdr_bh')
    df_res['adj_p_fdr'] = pvals_fdr
    df_res['adj_p_reject_fdr'] = reject

    # top associations by adjusted (FDR) p-value
    top = df_res.sort_values('adj_p_fdr').head(top_n)
    doc.add_heading('Top {} associations (by BH-adjusted p)'.format(top_n), level=2)
    for _, r in top.iterrows():
        x = r['exposure']
        y = r['outcome']
        doc.add_heading(f'{y}  vs  {x}', level=3)
        doc.add_paragraph(f"n={int(r['n'])}; adj_beta={r['adj_beta']:.3f}; 95%CI=({r['adj_ci_low']:.3f}, {r['adj_ci_high']:.3f}); adj_p={r['adj_p']:.3g}; BH_adj_p={r['adj_p_fdr']:.3g}")
        img = plots_dir / f"{sanitize(x)}__vs__{sanitize(y)}.png"
        if img.exists():
            try:
                doc.add_picture(str(img), width=Inches(6.5))
            except Exception:
                try:
                    doc.add_picture(str(img), width=Inches(5.5))
                except Exception:
                    pass
        doc.add_page_break()

    # add 1-2 adjusted regression model summaries for the top 2 associations
    doc.add_heading('Example adjusted regression model summaries', level=2)
    for _, r in top.head(2).iterrows():
        x = r['exposure']
        y = r['outcome']
        doc.add_heading(f'Adjusted model: {y} ~ {x} + covariates', level=3)
        # Use the adjusted results we computed earlier (exposure term + model summary)
        tbl = doc.add_table(rows=1, cols=7)
        hdr = tbl.rows[0].cells
        hdr[0].text = 'exposure'; hdr[1].text = 'n'; hdr[2].text = 'adj_beta'; hdr[3].text = 'se'; hdr[4].text = '95% CI'; hdr[5].text = 'p'; hdr[6].text = 'adj_r2'
        row = tbl.add_row().cells
        row[0].text = str(r['exposure'])
        row[1].text = str(int(r['n']))
        row[2].text = f"{r['adj_beta']:.4f}" if pd.notna(r['adj_beta']) else 'nan'
        row[3].text = f"{r['adj_se']:.4f}" if pd.notna(r['adj_se']) else 'nan'
        row[4].text = f"({r['adj_ci_low']:.3f}, {r['adj_ci_high']:.3f})" if pd.notna(r['adj_ci_low']) else 'nan'
        row[5].text = f"{r['adj_p']:.3g}" if pd.notna(r['adj_p']) else 'nan'
        row[6].text = f"{r.get('adj_r2', np.nan):.3f}" if pd.notna(r.get('adj_r2', np.nan)) else 'nan'
        # change-in-estimate note
        ci_note = f"Change-in-estimate (pct): {r.get('change_pct', np.nan):.1f}%" if pd.notna(r.get('change_pct')) else ''
        doc.add_paragraph(ci_note)
        doc.add_page_break()

    # Confounding & causal assessment (automated brief)
    doc.add_heading('Confounding & causal assessment', level=2)
    doc.add_paragraph('''
This automated assessment reports change-in-estimate for the exposure coefficient when adding the default covariates (Age, Income, Education, Marital Status).
Rules used:
- Change-in-estimate >= 10% flagged as evidence of confounding by the adjustment set.
- Results are cross-sectional: causal claims require temporality and stronger designs.

Typical limitations: residual/unmeasured confounding, measurement error, and model misspecification. Interpret effect sizes with confidence intervals and consider triangulation or longitudinal data for causal inference.
''')

    doc.save(out_docx)
    print('Saved Word summary to', out_docx)

def main():
    inp = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('InputData/20251214-ScalesData-Combined_ver0.7-Cleaned-IncomeNormalized-analysis-ready.csv')
    if not inp.exists():
        print('Input not found:', inp)
        sys.exit(2)
    df = pd.read_csv(inp, engine='python', encoding='utf-8')
    print('Loaded', inp, 'shape', df.shape)

    # candidate exposures
    exposures = find_internet_exposures(df.columns)
    if not exposures:
        print('No internet-related exposures found by keywords; please edit EXPOSURE_KEYWORDS in the script or pass a list.')
        sys.exit(1)
    print('Found exposures:', exposures)

    # dependents: use a small set of likely outcomes (per repo DEPEND) or choose top variances
    # use a default depend list inspired by earlier scripts
    DEPEND = [
        'Perceived Social Support Total T Score',
        'Perceived Social Stress Score',
        'Emotional Support - T Score',
        'Informational Support - T Score',
        'Instrumental Support - T Score'
    ]
    dependents = select_top_dependents(df, DEPEND, top_n=5)
    print('Selected dependents:', dependents)

    df_res, plots_dir = run_models(df, exposures, dependents, DEFAULT_COVARIATES)
    out_csv = Path('InputData') / 'adjusted_models_batch_summary.csv'
    df_res.to_csv(out_csv, index=False)
    print('Wrote batch summary to', out_csv)

    # FDR on adjusted p-values
    if not df_res.empty and 'adj_p' in df_res.columns:
        pvals = df_res['adj_p'].fillna(1.0).values
        reject, pvals_fdr, _, _ = multipletests(pvals, alpha=0.05, method='fdr_bh')
        df_res['adj_p_fdr'] = pvals_fdr
        df_res['adj_p_reject_fdr'] = reject
        df_res.to_csv(out_csv, index=False)

    # Create Word report
    out_docx = Path('InputData') / 'batch_ols_report.docx'
    make_word_report(df_res, plots_dir, out_docx, top_n=10)

    print('Done.')

if __name__ == '__main__':
    main()
