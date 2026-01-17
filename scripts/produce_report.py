#!/usr/bin/env python3
"""
Produce EDA plots and a PDF report for the ScalesData CSV.
Generates:
 - outputs/plots/<col>_hist.png and <col>_box.png for numeric columns
 - outputs/plots/<col>_bar.png for categorical columns
 - scales_plots_report.pdf (combined report)

Usage: python3 produce_report.py /path/to/CSV
"""
import sys
from pathlib import Path
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
from matplotlib.backends.backend_pdf import PdfPages
try:
    from docx import Document
    from docx.shared import Inches, Pt
except Exception:
    Document = None
    Inches = None
    Pt = None

# Ensure figures and saved output use a white background so text placed
# in the figure (e.g. footer/summary) remains visible in PDF viewers.
plt.rcParams.update({
    'figure.facecolor': 'white',
    'axes.facecolor': 'white',
    'savefig.facecolor': 'white'
})

# helper
def sanitize(s):
    return ''.join(c if c.isalnum() or c in (' ', '_', '-') else '_' for c in s).strip().replace(' ', '_')

CSV = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('/home/pranav/PhD-Jaya/InputData/20251214-ScalesData-Combined_ver0.7-Cleaned-IncomeNormalized-analysis-ready.csv')
OUT_DIR = CSV.parent / 'outputs' / 'plots'
OUT_DIR.mkdir(parents=True, exist_ok=True)
PDF_OUT = CSV.parent / 'scales_plots_report.pdf'

print('Reading', CSV)
df = pd.read_csv(CSV, engine='python', encoding='utf-8')
print('Shape:', df.shape)

# detect numeric and categorical
numeric = df.select_dtypes(include=[np.number]).columns.tolist()
cat = df.select_dtypes(include=['object','category']).columns.tolist()

# Flag numeric columns with IQR-based outliers or mean/median divergence
flagged_numeric = []
for col in numeric:
    ser = df[col].dropna()
    if ser.empty:
        continue
    q1 = ser.quantile(0.25)
    q3 = ser.quantile(0.75)
    iqr = q3 - q1
    lower = q1 - 1.5 * iqr
    upper = q3 + 1.5 * iqr
    outliers = ser[(ser < lower) | (ser > upper)]
    mean = ser.mean()
    median = ser.median()
    std = ser.std()
    pct_out = outliers.count() / ser.count() if ser.count() > 0 else 0
    if pct_out > 0.01 or abs(mean - median) > max(1.5 * std, 1e-9):
        flagged_numeric.append(col)

# For convenience, also include top N numeric columns if flagged list small
if len(flagged_numeric) < 8:
    # add top variance numeric columns
    variances = df[numeric].var().sort_values(ascending=False)
    for c in variances.index[:min(8, len(variances))]:
        if c not in flagged_numeric:
            flagged_numeric.append(c)

print('Flagged numeric columns:', flagged_numeric)

# Prepare PDF
pp = PdfPages(str(PDF_OUT))
# Title page (explicitly use a Figure so we can control facecolor)
fig = plt.figure(figsize=(11.7,8.3), facecolor='white')
ax = fig.add_subplot(111)
ax.axis('off')
fig.text(0.5, 0.7, 'Scales Data EDA Report', ha='center', va='center', fontsize=22)
fig.text(0.5, 0.6, f'File: {CSV.name}', ha='center', va='center', fontsize=10)
fig.text(0.5, 0.55, f'Rows: {len(df)}  Columns: {len(df.columns)}', ha='center', va='center', fontsize=10)
fig.text(0.05, 0.3, 'Notes:\n - Numeric flagged by simple IQR/outlier heuristics\n - Categorical plots limited to top categories', fontsize=9)
pp.savefig(fig, bbox_inches='tight', facecolor=fig.get_facecolor())
plt.close(fig)
# Plot every column as its own page in the PDF
import textwrap

all_columns = list(df.columns)
print('Plotting all columns (count):', len(all_columns))

for col in all_columns:
    ser = df[col]
    # handle numeric columns: histogram + boxplot on one page
    if pd.api.types.is_numeric_dtype(ser):
        ser_clean = ser.dropna()
        n = int(ser_clean.count()) if not ser_clean.empty else 0
        mean = float(ser_clean.mean()) if n>0 else np.nan
        median = float(ser_clean.median()) if n>0 else np.nan
        std = float(ser_clean.std()) if n>0 else np.nan
        mn = float(ser_clean.min()) if n>0 else np.nan
        mx = float(ser_clean.max()) if n>0 else np.nan
        q1 = float(ser_clean.quantile(0.25)) if n>0 else np.nan
        q3 = float(ser_clean.quantile(0.75)) if n>0 else np.nan
        iqr = q3 - q1 if n>0 else np.nan
        lower = q1 - 1.5 * iqr if n>0 else np.nan
        upper = q3 + 1.5 * iqr if n>0 else np.nan
        outliers = ser_clean[(ser_clean < lower) | (ser_clean > upper)] if n>0 else ser_clean.iloc[0:0]
        n_out = int(outliers.count())
        pct_out = n_out / n if n>0 else 0
        skewness = float(ser_clean.skew()) if n>2 else 0.0
        skew_note = ''
        if skewness > 0.8:
            skew_note = 'Strong right skew (consider log or Box-Cox transform).'
        elif skewness > 0.3:
            skew_note = 'Moderate right skew (consider transform).'
        elif skewness < -0.8:
            skew_note = 'Strong left skew (consider reflection + transform).'
        elif skewness < -0.3:
            skew_note = 'Moderate left skew.'

        if pct_out >= 0.05:
            outlier_sugg = 'High proportion of outliers — consider inspection, winsorizing, or robust methods.'
        elif pct_out >= 0.01:
            outlier_sugg = 'Some outliers present — consider trimming or robust estimators.'
        else:
            outlier_sugg = 'Few outliers — standard methods likely OK.'

        # prepare formatted numeric strings (handle NaN)
        mean_f = f"{mean:.2f}" if not np.isnan(mean) else "nan"
        median_f = f"{median:.2f}" if not np.isnan(median) else "nan"
        std_f = f"{std:.2f}" if not np.isnan(std) else "nan"
        mn_f = f"{mn:.2f}" if not np.isnan(mn) else "nan"
        mx_f = f"{mx:.2f}" if not np.isnan(mx) else "nan"
        skew_f = f"{skewness:.2f}"

        desc = (
            f'Count={n}; mean={mean_f}; median={median_f}; std={std_f}; '
            f'min={mn_f}; max={mx_f}; outliers={n_out} ({pct_out:.1%}). '
            f'Skewness={skew_f}. {skew_note} {outlier_sugg}'
        )

        fig, axes = plt.subplots(nrows=2, ncols=1, figsize=(8.5,11), facecolor='white', gridspec_kw={'height_ratios':[3,1]})
        # histogram (top)
        ax_hist = axes[0]
        sns.histplot(ser_clean, kde=True, ax=ax_hist)
        ax_hist.set_title(f'Histogram: {col}')
        ax_hist.set_xlabel(col)
        # boxplot (bottom)
        ax_box = axes[1]
        sns.boxplot(x=ser_clean, orient='h', ax=ax_box)
        ax_box.set_title(f'Boxplot: {col}')
        ax_box.set_xlabel('')

        # description footer
        wrapped = textwrap.fill(desc, 200)
        fig.patch.set_facecolor('white')
        fig.text(0.01, 0.02, wrapped, ha='left', va='bottom', fontsize=9)
        plt.tight_layout(rect=[0,0.06,1,0.98])

        # save individual images (optional, keep old filenames)
        fname_hist = OUT_DIR / f'{sanitize(col)}_hist.png'
        fname_box = OUT_DIR / f'{sanitize(col)}_box.png'
        # save the combined page to PDF and also save the separate files
        # save high-resolution images for legibility in PDF and Word
        fig.savefig(fname_hist, bbox_inches='tight', facecolor=fig.get_facecolor(), dpi=200)
        # also save a box-only image for consistency
        # create a separate small fig for the boxplot image
        fig_box_small, axb = plt.subplots(figsize=(8,3), facecolor='white')
        sns.boxplot(x=ser_clean, orient='h', ax=axb)
        axb.set_title(f'Boxplot: {col}')
        plt.tight_layout()
        fig_box_small.savefig(fname_box, bbox_inches='tight', facecolor=fig_box_small.get_facecolor(), dpi=200)
        plt.close(fig_box_small)

        pp.savefig(fig, bbox_inches='tight', facecolor=fig.get_facecolor())
        plt.close(fig)

    else:
        # categorical: bar chart per variable
        ser_cat = ser.fillna('<<MISSING>>').astype(str)
        vc = ser_cat.value_counts().head(50)
        fig = plt.figure(figsize=(8.5,11), facecolor='white')
        ax = fig.add_subplot(111)
        sns.barplot(y=vc.index.astype(str), x=vc.values, palette='viridis', ax=ax)
        ax.set_title(f'Value counts: {col}')
        ax.set_xlabel('Count')
        ax.set_ylabel(col)
        total = int(ser_cat.shape[0])
        lines = []
        for k, v in vc.items():
            pct = v / total if total>0 else 0
            lines.append(f'{k}: {v} ({pct:.1%})')
        top_pct = vc.iloc[0] / total if total>0 and len(vc)>0 else 0
        if top_pct > 0.8:
            importance = 'Single category dominates — low variability.'
        elif top_pct > 0.4:
            importance = 'Top category is sizeable — consider grouping smaller categories.'
        else:
            importance = 'Healthy spread across categories.'
        desc_full = 'Top categories: ' + ' | '.join(lines) + f'. Note: {importance}'
        fig.patch.set_facecolor('white')
        fig.text(0.01, 0.02, '\n'.join(textwrap.wrap(desc_full, 300)), ha='left', va='bottom', fontsize=9)
        plt.tight_layout(rect=[0,0.06,1,0.98])
        fname_bar = OUT_DIR / f'{sanitize(col)}_bar.png'
        # save high-resolution bar image
        fig.savefig(fname_bar, bbox_inches='tight', facecolor=fig.get_facecolor(), dpi=200)
        pp.savefig(fig, bbox_inches='tight', facecolor=fig.get_facecolor())
        plt.close(fig)
pp.close()
print('Saved PDF report to', PDF_OUT)
print('Saved individual plots to', OUT_DIR)

# (sanitize defined above)

# Write small index of files
with open(CSV.parent / 'outputs' / 'plots' / 'index.txt', 'w', encoding='utf-8') as fh:
    for p in sorted(OUT_DIR.iterdir()):
        fh.write(str(p.name) + '\n')

print('Done.')

# Additionally create a Microsoft Word (.docx) report with the same content as the PDF
DOCX_OUT = CSV.parent / 'scales_plots_report.docx'
if Document is None:
    print('python-docx not available; skipping Word document creation. To enable, pip install python-docx')
else:
    print('Creating Word document at', DOCX_OUT)
    doc = Document()
    # Title page
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    h = doc.add_heading('Scales Data EDA Report', level=0)
    h.alignment = 1
    doc.add_paragraph(f'File: {CSV.name}')
    doc.add_paragraph(f'Rows: {len(df)}  Columns: {len(df.columns)}')
    doc.add_paragraph('Notes:')
    doc.add_paragraph(' - Numeric flagged by simple IQR/outlier heuristics')
    doc.add_paragraph(' - Categorical plots limited to top categories')
    doc.add_page_break()

    # Add one section per column mirroring the PDF
    for col in all_columns:
        doc.add_heading(str(col), level=2)
        ser = df[col]
        if pd.api.types.is_numeric_dtype(ser):
            hist_img = OUT_DIR / f'{sanitize(col)}_hist.png'
            box_img = OUT_DIR / f'{sanitize(col)}_box.png'
            if hist_img.exists():
                try:
                    doc.add_picture(str(hist_img), width=Inches(6.5))
                except Exception:
                    pass
            # add boxplot smaller
            if box_img.exists():
                try:
                    doc.add_picture(str(box_img), width=Inches(6.5))
                except Exception:
                    pass
            # add the same description footer as text
            ser_clean = ser.dropna()
            n = int(ser_clean.count()) if not ser_clean.empty else 0
            mean = float(ser_clean.mean()) if n>0 else np.nan
            median = float(ser_clean.median()) if n>0 else np.nan
            std = float(ser_clean.std()) if n>0 else np.nan
            mn = float(ser_clean.min()) if n>0 else np.nan
            mx = float(ser_clean.max()) if n>0 else np.nan
            q1 = float(ser_clean.quantile(0.25)) if n>0 else np.nan
            q3 = float(ser_clean.quantile(0.75)) if n>0 else np.nan
            iqr = q3 - q1 if n>0 else np.nan
            lower = q1 - 1.5 * iqr if n>0 else np.nan
            upper = q3 + 1.5 * iqr if n>0 else np.nan
            outliers = ser_clean[(ser_clean < lower) | (ser_clean > upper)] if n>0 else ser_clean.iloc[0:0]
            n_out = int(outliers.count())
            pct_out = n_out / n if n>0 else 0
            skewness = float(ser_clean.skew()) if n>2 else 0.0
            mean_f = f"{mean:.2f}" if not np.isnan(mean) else "nan"
            median_f = f"{median:.2f}" if not np.isnan(median) else "nan"
            std_f = f"{std:.2f}" if not np.isnan(std) else "nan"
            mn_f = f"{mn:.2f}" if not np.isnan(mn) else "nan"
            mx_f = f"{mx:.2f}" if not np.isnan(mx) else "nan"
            desc = (
                f'Count={n}; mean={mean_f}; median={median_f}; std={std_f}; '
                f'min={mn_f}; max={mx_f}; outliers={n_out} ({pct_out:.1%}). Skewness={skewness:.2f}.'
            )
            p = doc.add_paragraph()
            run = p.add_run(desc)
            run.font.size = Pt(10)
        else:
            bar_img = OUT_DIR / f'{sanitize(col)}_bar.png'
            if bar_img.exists():
                try:
                    doc.add_picture(str(bar_img), width=Inches(6.5))
                except Exception:
                    pass
            # add top categories text
            ser_cat = ser.fillna('<<MISSING>>').astype(str)
            vc = ser_cat.value_counts().head(50)
            total = int(ser_cat.shape[0])
            lines = []
            for k, v in vc.items():
                pct = v / total if total>0 else 0
                lines.append(f'{k}: {v} ({pct:.1%})')
            desc_full = 'Top categories: ' + ' | '.join(lines)
            p = doc.add_paragraph()
            run = p.add_run(desc_full)
            run.font.size = Pt(10)
        doc.add_page_break()
    try:
        doc.save(str(DOCX_OUT))
        print('Saved Word report to', DOCX_OUT)
    except Exception as e:
        print('Could not save Word doc:', e)
