"""
relationship_analysis_internet.py

Performs pairwise relationship analysis (correlation, ANOVA, chi-square) between specified internet-related independent variables and a set of dependent variables.

Outputs:
- relation_summary.csv: Raw pairwise test results
- relation_summary_with_fdr.csv: FDR-corrected results
- relation_report.pdf: Plots and summary (PDF)
- relation_report.docx: Plots and summary (Word)

Usage:
    python3 scripts/relationship_analysis_internet.py InputData/20251214-ScalesData-Combined_ver0.7-Cleaned-IncomeNormalized.csv

Dependencies:
- pandas, numpy, matplotlib, seaborn, scipy, statsmodels, python-docx, fpdf
"""
import sys
from pathlib import Path
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
from scipy import stats
from statsmodels.stats.multitest import multipletests

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    Document = None
try:
    from fpdf import FPDF
except ImportError:
    FPDF = None

# List of independent (internet-related) variables
INDEPENDENTS = [
    'Functional Impairment - Raw Score',
    'Withdrawl - Raw Score',
    'Occupational & Relationship Consequences - Raw Score',
    'Compulsive Behaviour - Raw Score',
    'Obsession with Internet - Raw Score',
    'Internet as a Source of Recreation - Raw Score',
    'Enhanced Socialization - Raw Score',
    'Perceived Control of Internet Use - Raw Score',
    'Functional Impairment - Weighted Score',
    'Withdrawl - Weighted Score',
    'Occupational & Relationship Consequences - Weighted Score',
    'Compulsive Behaviour - Weighted Score',
    'Obsession with Internet - Weighted Score',
    'Internet as a Source of Recreation - Weighted Score',
    'Enhanced Socialization - Weighted Score',
    'Perceived Control of Internet Use - Weighted Score',
    'Internet Usage Total Score (Weighted)',
    'Functional Impairment  - Category',
    'Withdrawl – Category',
    'Occupational & Relationship Consequences – Category',
    'Compulsive Behaviour – Category',
    'Obsession with Internet – Category',
    'Internet as a Source of Recreation – Category',
    'Enhanced Socialization - Category',
    'Perceived Control of Internet Use - Category',
    'Internet Usage Total Categoty',
]

# List of dependent variables
DEPENDENTS = [
    'Emotional Support - Raw Score',
    'Informational Support - Raw Score',
    'Instrumental Support - Raw Score',
    'Perceived Social Support Total Raw Score',
    'Emotional Support - T Score',
    'Informational Support - T Score',
    'Instrumental Support - T Score',
    'Perceived Social Support Total T Score',
    'Emotional Support - Rank',
    'Informational Support - Rank',
    'Instrumental Support - Rank',
    'Perceived Social Support TOTAL - Rank',
    'Power Self Concept - Real Self Concept Score',
    'Power Self Concept - Ideal Self Concept score',
    'Power Self Concept – Social Self Concept Score',
    'Social Self Concept - Real Self Concept Score',
    'Social Self Concept – Ideal Self Concept Score',
    'Social Self Concept – Social Self Concept Score',
    'Ability Self Concept - Real Self Concept Score',
    'Ability Self Concept - Ideal Self Concept Score',
    'Ability Self Concept – Social Self Concept Score',
    'Physical Self Concept - Real Self Concept Score',
    'Physical Self Concept – Ideal Self Concept Score',
    'Physical Self Concept – Social Self Concept Score',
    'Psychological Self Concept - Real Self Concept Score',
    'Psychological Self Concept – Ideal Self Concept Score',
    'Psychological Self Concept – Social Self Concept Score',
    'Real Self Concept Total Score',
    'Ideal Self Concept Total Score',
    'Social Self Concept Total Score',
    'Real Self Concept Total Category',
    'Ideal Self Concept Total Category',
    'Social Self Concept Total Category',
    'Perceived Social Stress Score',
    'Perceived Social Stress Category',
]

def main():
    if len(sys.argv) < 2:
        print("Usage: python3 scripts/relationship_analysis_internet.py <input_csv>")
        sys.exit(1)
    csv_path = Path(sys.argv[1])
    df = pd.read_csv(csv_path)
    # Ensure Marital Status column exists and drop missing
    if 'Marital Status' not in df.columns:
        print('Marital Status column not found in input data.')
        sys.exit(1)
    df = df.dropna(subset=['Marital Status'])
    groups = {'Married': 'married', 'Single': 'single'}
    for group_label, group_key in groups.items():
        group_mask = df['Marital Status'].astype(str).str.lower().str.contains(group_key)
        group_df = df[group_mask].copy()
        if group_df.empty:
            print(f'No data for {group_label} females. Skipping.')
            continue
        # Prepare output directories and filenames
        outdir = Path(f'outputs_internet/{group_key}')
        out_plots = outdir / 'plots'
        # Delete previous outputs if they exist
        import shutil, os
        if outdir.exists():
            shutil.rmtree(outdir)
        out_plots.mkdir(parents=True, exist_ok=True)
        pdf_out = outdir / 'relation_report.pdf'
        csv_out = outdir / 'relation_summary.csv'
        fdr_csv_out = outdir / 'relation_summary_with_fdr.csv'
        docx_out = outdir / 'relation_report.docx'
        from matplotlib.backends.backend_pdf import PdfPages
        if Document is not None:
            doc = Document()
            doc.add_heading(f'Internet Variable Relation Report: {group_label} Females', level=1)
            doc.add_paragraph(f'Source file: {csv_path.name}')
        pp = PdfPages(str(pdf_out))
        results = []
        for indep in INDEPENDENTS:
            if indep not in group_df.columns:
                continue
            for dep in DEPENDENTS:
                if dep not in group_df.columns:
                    continue
                x = group_df[indep]
                y = group_df[dep]
                res = {'independent': indep, 'dependent': dep, 'test': None, 'statistic': None, 'pvalue': None, 'notes': ''}
                fig = plt.figure(figsize=(8.5, 11), facecolor='white')
                ax = fig.add_subplot(111)
                try:
                    if pd.api.types.is_numeric_dtype(x) and pd.api.types.is_numeric_dtype(y):
                        r, p = stats.spearmanr(x, y, nan_policy='omit')
                        res['test'] = 'spearman'
                        res['statistic'] = r
                        res['pvalue'] = p
                        sns.regplot(x=indep, y=dep, data=group_df, ax=ax, scatter_kws={'s':10}, line_kws={'color':'red'})
                        ax.set_title(f'{dep} vs {indep} (Spearman r={r:.2f}, p={p:.2g})')
                    elif isinstance(x.dtype, pd.CategoricalDtype) or x.dtype == object:
                        if pd.api.types.is_numeric_dtype(y):
                            groups_ = [y[x == cat] for cat in pd.unique(x.dropna())]
                            if len(groups_) > 1:
                                f, p = stats.f_oneway(*groups_)
                                res['test'] = 'anova'
                                res['statistic'] = f
                                res['pvalue'] = p
                                sns.boxplot(x=indep, y=dep, data=group_df, ax=ax)
                                ax.set_title(f'{dep} by {indep} (ANOVA F={f:.2f}, p={p:.2g})')
                        else:
                            ct = pd.crosstab(x, y)
                            if ct.shape[0] > 1 and ct.shape[1] > 1:
                                chi2, p, _, _ = stats.chi2_contingency(ct)
                                res['test'] = 'chi2'
                                res['statistic'] = chi2
                                res['pvalue'] = p
                                sns.heatmap(ct, annot=True, fmt='d', cmap='Blues', ax=ax)
                                ax.set_title(f'Cross-tab: {indep} x {dep} (chi2={chi2:.2f}, p={p:.2g})')
                    else:
                        res['notes'] = 'Test not run: unsupported type combination.'
                except Exception as e:
                    res['notes'] = f'Plot/test error: {e}'
                # Save plot
                fname = out_plots / f"{indep.replace(' ', '_').replace('/', '_')}_vs_{dep.replace(' ', '_').replace('/', '_')}.png"
                fig.savefig(fname, dpi=200, bbox_inches='tight', facecolor=fig.get_facecolor())
                plt.close(fig)
                # Add to PDF
                try:
                    img = plt.imread(str(fname))
                    fig2 = plt.figure(figsize=(8.5, 11), facecolor='white')
                    ax2 = fig2.add_subplot(111)
                    ax2.axis('off')
                    ax2.imshow(img)
                    pp.savefig(fig2, bbox_inches='tight', facecolor=fig2.get_facecolor())
                    plt.close(fig2)
                except Exception:
                    pass
                # Add to DOCX
                if Document is not None:
                    doc.add_heading(f'{dep} vs {indep}', level=2)
                    summary_para = f"Test: {res['test']}; statistic: {res['statistic']}; p-value: {res['pvalue']}; {res['notes']}"
                    doc.add_paragraph(summary_para)
                    try:
                        doc.add_picture(str(fname), width=Inches(6.5))
                    except Exception:
                        pass
                    doc.add_page_break()
                results.append(res)
        pp.close()
        # Save raw results
        df_res = pd.DataFrame(results)
        df_res.to_csv(csv_out, index=False)
        # FDR correction
        if not df_res.empty:
            reject, pvals_fdr, _, _ = multipletests(df_res['pvalue'], method='fdr_bh')
            df_res['pvalue_fdr'] = pvals_fdr
            df_res['significant_fdr'] = reject
            df_res.to_csv(fdr_csv_out, index=False)
        # Save DOCX
        if Document is not None:
            doc.save(docx_out)
        print(f'[{group_label}] Wrote summary to {csv_out}, plots to {out_plots}, PDF to {pdf_out}')

if __name__ == '__main__':
    main()
