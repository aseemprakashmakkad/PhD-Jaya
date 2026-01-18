"""
stratified_statistical_analysis.py

Performs statistical analysis (Spearman, ANOVA, Chi-square) between internet-related exposures, demographic covariates, and dependent variables, stratified by marital status (married vs. unmarried). Summarizes key findings and generates a well-formatted MS Word document.

Usage:
    python3 scripts/stratified_statistical_analysis.py relation_summary_with_fdr.csv InputData/20251214-ScalesData-Combined_ver0.7-Cleaned-IncomeNormalized.csv

Outputs:
    outputs/stratified_statistical_analysis.docx

Dependencies:
    pandas, numpy, python-docx
"""
import sys
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def summarize_spearman(df):
    rows = []
    for _, row in df[df['test'].str.contains('spearman', na=False)].iterrows():
        r = row['statistic']
        p = row['pvalue']
        indep = row['independent']
        dep = row['dependent']
        if pd.notnull(r) and pd.notnull(p):
            if abs(r) >= 0.5:
                strength = 'strong'
            elif abs(r) >= 0.3:
                strength = 'moderate'
            elif abs(r) >= 0.1:
                strength = 'weak'
            else:
                strength = 'very weak'
            direction = 'positive' if r > 0 else 'negative'
            sig = 'statistically significant' if p < 0.05 else 'not significant'
            comment = f"{indep} vs. {dep}: {strength} {direction} association (r={r:.2f}, p={p:.3g}, {sig})"
            rows.append(comment)
    return rows

def summarize_anova(df):
    rows = []
    for _, row in df[df['test'].str.contains('anova', na=False)].iterrows():
        F = row['statistic']
        p = row['pvalue']
        indep = row['independent']
        dep = row['dependent']
        if pd.notnull(F) and pd.notnull(p):
            strength = 'strong' if F >= 5 else 'moderate' if F >= 2 else 'weak'
            sig = 'statistically significant' if p < 0.05 else 'not significant'
            comment = f"{indep} vs. {dep}: {strength} group differences (F={F:.2f}, p={p:.3g}, {sig})"
            rows.append(comment)
    return rows

def summarize_chi2(df):
    rows = []
    for _, row in df[df['test'].str.contains('chi2', na=False)].iterrows():
        chi2 = row['statistic']
        p = row['pvalue']
        indep = row['independent']
        dep = row['dependent']
        if pd.notnull(chi2) and pd.notnull(p):
            strength = 'strong' if chi2 >= 10 else 'moderate' if chi2 >= 5 else 'weak'
            sig = 'statistically significant' if p < 0.05 else 'not significant'
            comment = f"{indep} vs. {dep}: {strength} non-random association (χ²={chi2:.2f}, p={p:.3g}, {sig})"
            rows.append(comment)
    return rows

def main():
    if len(sys.argv) < 3:
        print("Usage: python3 scripts/stratified_statistical_analysis.py relation_summary_with_fdr.csv InputData/20251214-ScalesData-Combined_ver0.7-Cleaned-IncomeNormalized.csv")
        sys.exit(1)
    summary_csv = Path(sys.argv[1])
    data_csv = Path(sys.argv[2])
    df_summary = pd.read_csv(summary_csv)
    df_data = pd.read_csv(data_csv)
    # Drop rows with missing Marital Status
    if 'Marital Status' in df_data.columns:
        df_data = df_data.dropna(subset=['Marital Status'])
    doc = Document()
    doc.add_heading("Stratified Statistical Analysis by Marital Status", 0)
    for group in ['Married', 'Unmarried']:
        doc.add_heading(f"{group} Females", 1)
        # Filter data for group, robust to missing values
        if 'Marital Status' not in df_data.columns:
            doc.add_paragraph("Marital Status column not found in data.")
            continue
        # Only use rows where Marital Status is not null and matches group
        group_mask = df_data['Marital Status'].astype(str).str.lower().str.contains(group.lower(), na=False)
        group_df = df_data[group_mask].copy()
        if group_df.empty:
            doc.add_paragraph(f"No data for {group} females.")
            continue
        # Recompute stats for this group
        results = []
        for _, row in df_summary.iterrows():
            indep = row['independent']
            dep = row['dependent']
            test = row['test']
            if indep not in group_df.columns or dep not in group_df.columns:
                continue
            x = group_df[indep]
            y = group_df[dep]
            # Drop rows with missing values in either variable
            valid = x.notna() & y.notna()
            x_valid = x[valid]
            y_valid = y[valid]
            if len(x_valid) == 0 or len(y_valid) == 0:
                continue
            if test == 'spearman':
                from scipy.stats import spearmanr
                r, p = spearmanr(x_valid, y_valid, nan_policy='omit')
                if pd.notnull(r) and pd.notnull(p):
                    if abs(r) >= 0.5:
                        strength = 'strong'
                    elif abs(r) >= 0.3:
                        strength = 'moderate'
                    elif abs(r) >= 0.1:
                        strength = 'weak'
                    else:
                        strength = 'very weak'
                    direction = 'positive' if r > 0 else 'negative'
                    sig = 'statistically significant' if p < 0.05 else 'not significant'
                    comment = f"{indep} vs. {dep}: {strength} {direction} association (r={r:.2f}, p={p:.3g}, {sig})"
                    results.append(comment)
            elif test == 'anova':
                from scipy.stats import f_oneway
                if x_valid.nunique() > 1:
                    groups = [y_valid[x_valid == cat] for cat in pd.unique(x_valid)]
                    groups = [g for g in groups if len(g) > 0]
                    if len(groups) > 1:
                        F, p = f_oneway(*groups)
                        strength = 'strong' if F >= 5 else 'moderate' if F >= 2 else 'weak'
                        sig = 'statistically significant' if p < 0.05 else 'not significant'
                        comment = f"{indep} vs. {dep}: {strength} group differences (F={F:.2f}, p={p:.3g}, {sig})"
                        results.append(comment)
            elif test == 'chi2':
                from scipy.stats import chi2_contingency
                ct = pd.crosstab(x_valid, y_valid)
                if ct.shape[0] > 1 and ct.shape[1] > 1:
                    chi2, p, _, _ = chi2_contingency(ct)
                    strength = 'strong' if chi2 >= 10 else 'moderate' if chi2 >= 5 else 'weak'
                    sig = 'statistically significant' if p < 0.05 else 'not significant'
                    comment = f"{indep} vs. {dep}: {strength} non-random association (χ²={chi2:.2f}, p={p:.3g}, {sig})"
                    results.append(comment)
        if results:
            for c in results:
                doc.add_paragraph(c, style='List Bullet')
        else:
            doc.add_paragraph("No significant results for this group.")
    doc.save('outputs/stratified_statistical_analysis.docx')
    print('Saved stratified analysis to outputs/stratified_statistical_analysis.docx')

if __name__ == '__main__':
    main()
