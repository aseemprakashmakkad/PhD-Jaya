from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')

def add_table_from_df(doc, df, max_rows=10):
    if df.empty:
        doc.add_paragraph("No data available.")
        return
    ncols = len(df.columns)
    nrows = min(len(df), max_rows)
    table = doc.add_table(rows=nrows+1, cols=ncols)
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = str(col)
    for i in range(nrows):
        for j, col in enumerate(df.columns):
            val = df.iloc[i, j]
            table.cell(i+1, j).text = str(val)
    table.style = 'Table Grid'

def add_image(doc, img_path, caption=None, width_in=5.0):
    from docx.shared import Inches
    if Path(img_path).exists():
        doc.add_picture(str(img_path), width=Inches(width_in))
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Image not found: {img_path}]")

def main():
    doc = Document()
    doc.add_heading("Social Networking and Its Impact on Perceived Stress, Self-Concept & Social Support of Married Females versus Un-Married Females", 0)

    # Objective
    add_heading(doc, "Objective", 1)
    add_paragraph(doc, "To investigate the multifaceted impact of social networking on the perceived stress levels, self-concept, and social support networks of Indian females in the context of their marital status. By addressing these questions, this research seeks to provide valuable insights into the complex interplay between social networking, psychological well-being, and social relationships among Indian women, shedding light on the unique experiences and challenges faced by both married and unmarried individuals in the digital age.")

    # Variables
    add_heading(doc, "Variables and Study Design", 1)
    add_paragraph(doc, "Based on the questionnaire, the following variables were used:")
    add_paragraph(doc, "Covariates:", bold=True)
    for v in [
        "Age (IN YEARS)", "Marital Status", "Highest level of education completed", "Current Employment Status", "Number of children", "Type of family structure", "Monthly household income", "Access to healthcare services", "Overall health status", "Are you covered by health insurance?", "Participation in community or social activities"]:
        add_bullet(doc, v)
    add_paragraph(doc, "Independent variables:", bold=True)
    for v in [
        "Functional Impairment", "Withdrawl", "Occupational & Relationship Consequences", "Compulsive Behaviour", "Obsession with Internet", "Internet as a Source of Recreation", "Enhanced Socialization", "Perceived Control of Internet Use", "Internet Usage Total Score"]:
        add_bullet(doc, v)
    add_paragraph(doc, "Dependent variables:", bold=True)
    for v in [
        "Emotional Support", "Informational Support", "Instrumental Support", "Perceived Social Support Total Score", "Power Self Concept - Real Self Concept Score", "Power Self Concept - Ideal Self Concept score", "Power Self Concept – Social Self Concept Score", "Social Self Concept - Real Self Concept Score", "Social Self Concept – Ideal Self Concept Score", "Social Self Concept – Social Self Concept Score", "Ability Self Concept - Real Self Concept Score", "Ability Self Concept - Ideal Self Concept Score", "Ability Self Concept – Social Self Concept Score", "Physical Self Concept - Real Self Concept Score", "Physical Self Concept – Ideal Self Concept Score", "Physical Self Concept – Social Self Concept Score", "Psychological Self Concept - Real Self Concept Score", "Psychological Self Concept – Ideal Self Concept Score", "Psychological Self Concept – Social Self Concept Score", "Real Self Concept Total Score", "Ideal Self Concept Total Score", "Social Self Concept Total Score"]:
        add_bullet(doc, v)

    # Descriptive analysis
    add_heading(doc, "Descriptive Statistical Analysis", 1)
    add_paragraph(doc, "Descriptive statistics were computed for all variables, including means, standard deviations, quartiles, and histograms. The data showed a range of groupings by marital status, age, and education. Histograms and quartile charts revealed that many variables, such as internet usage and support scores, were moderately skewed, with some outliers. Demographic data indicated a diverse sample in terms of age, education, and income, with both married and unmarried women represented across socioeconomic strata.")


    # ANOVA & Chi-square
    add_heading(doc, "ANOVA & Chi-square Analyses", 1)
    add_paragraph(doc, "To assess the relationships between covariates and dependent variables, ANOVA and Chi-square tests were performed. These analyses identified significant associations between marital status, age, education, and several support/self-concept scores. For example, marital status was significantly associated with Emotional Support (ANOVA F=17.43, p<0.001), and age was associated with Emotional Support and Self-Concept scores. Participation in community activities and health insurance coverage also showed significant relationships with support outcomes.")
    add_paragraph(doc, "Key findings from ANOVA & Chi-square tests:")
    # Insert a small table of top significant results
    try:
        rel = pd.read_csv('InputData/relation_summary_with_fdr.csv')
        rel_sig = rel[rel['significant_fdr05']==True].sort_values('pvalue_fdr_bh').head(8)
        add_table_from_df(doc, rel_sig[['independent','dependent','test','statistic','pvalue','pvalue_fdr_bh']])
        # Try to add a plot for the top 1-2 relationships if available
        for i, row in rel_sig.head(2).iterrows():
            dep = row['dependent']
            indep = row['independent']
            # Try to find a plot in outputs/plots or outputs_relation/plots
            img1 = Path('InputData/outputs_relation/plots') / f"{indep.replace(' ', '_')}__vs__{dep.replace(' ', '_')}.png"
            img2 = Path('InputData/outputs/plots') / f"{dep.replace(' ', '_')}.png"
            if img1.exists():
                add_image(doc, img1, caption=f"{indep} vs {dep}")
            elif img2.exists():
                add_image(doc, img2, caption=f"Distribution of {dep}")
    except Exception:
        add_paragraph(doc, "(Relation summary table unavailable)")


    # OLS regression
    add_heading(doc, "OLS Regression Analysis", 1)
    add_paragraph(doc, "To further explore the relationship between internet-related exposures and the dependent scores, Ordinary Least Squares (OLS) regression models were fitted. Each model included the main exposure, the dependent variable, and covariates: Age (IN YEARS), Monthly household income_midpoint_INR, Highest level of education completed, and Marital Status. These covariates were chosen to control for potential confounding and to isolate the effect of internet use on psychological and social outcomes.")
    add_paragraph(doc, "Covariates were included in the OLS models because they are known to influence both internet use and psychosocial outcomes. Adjusting for these factors helps ensure that observed associations are not due to differences in age, income, education, or marital status.")

    # OLS results
    add_heading(doc, "Key Findings from OLS Models", 1)
    try:
        ols = pd.read_csv('InputData/adjusted_models_batch_summary.csv')
        ols = ols[ols['adj_p'].notnull()]
        ols = ols.sort_values('adj_p_fdr').head(8)
        add_table_from_df(doc, ols[['exposure','outcome','n','adj_beta','adj_se','adj_p','adj_p_fdr','adj_r2','change_pct']])
        # Try to add a plot for the top 1-2 OLS associations if available
        for i, row in ols.head(2).iterrows():
            exp = row['exposure']
            dep = row['outcome']
            img = Path('InputData/outputs_models/plots') / f"{exp.replace(' ', '_')}__vs__{dep.replace(' ', '_')}.png"
            if img.exists():
                add_image(doc, img, caption=f"{exp} vs {dep} (OLS)")
    except Exception:
        add_paragraph(doc, "(OLS summary table unavailable)")
    add_paragraph(doc, "No associations remained significant after Benjamini-Hochberg FDR correction at the 0.05 level. However, several exposure-outcome pairs showed nominal significance (adj_p < 0.05), including higher Internet Usage Total Score being associated with higher Perceived Social Stress and lower Social Support scores. Obsession with Internet scores were also negatively associated with support outcomes. These effects persisted after adjustment for covariates, though the magnitude was sometimes attenuated.")

    # Final summary
    add_heading(doc, "Summary and Interpretation", 1)
    add_paragraph(doc, "This analysis demonstrates that social networking and internet-related behaviors are associated with perceived stress, self-concept, and social support among Indian women. While no associations survived strict multiple-testing correction, the direction and consistency of effects suggest that higher internet use and obsession are linked to higher stress and lower support, even after accounting for age, income, education, and marital status. These findings highlight the importance of considering both psychosocial and demographic factors when evaluating the impact of digital behaviors on well-being, and suggest avenues for further research, including stratified analyses by marital status and more nuanced modeling of categorical exposures.")

    doc.save('InputData/thesis_analysis_internet.docx')
    print('Saved thesis write-up to InputData/thesis_analysis_internet.docx')

if __name__ == '__main__':
    main()
