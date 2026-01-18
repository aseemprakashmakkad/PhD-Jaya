from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')

def add_table_from_df(doc, df, max_rows=10):
    if df.empty:
        doc.add_paragraph("No data available.")
        return
    ncols = len(df.columns)
    nrows = min(len(df), max_rows)
    table = doc.add_table(rows=nrows+1, cols=ncols)
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = str(col)
    for i in range(nrows):
        for j, col in enumerate(df.columns):
            val = df.iloc[i, j]
            table.cell(i+1, j).text = str(val)
    table.style = 'Table Grid'

def main():
    doc = Document()
    doc.add_heading("Social Networking and Its Impact on Perceived Stress, Self-Concept & Social Support of Married Females versus Un-Married Females", 0)

    # Objective
    add_heading(doc, "Objective", 1)
    add_paragraph(doc, "To investigate the multifaceted impact of social networking on the perceived stress levels, self-concept, and social support networks of Indian females in the context of their marital status. By addressing these questions, this research seeks to provide valuable insights into the complex interplay between social networking, psychological well-being, and social relationships among Indian women, shedding light on the unique experiences and challenges faced by both married and unmarried individuals in the digital age.")

    # Scales used
    add_heading(doc, "Scales (Questionnaires) Used", 1)
    add_bullet(doc, "Internet Overuse Scale By Dr. Darshna Shah & Prof. Urmi Nanda Biswas")
    add_bullet(doc, "Self CONCEPT Scale by R. K. Saraswat, Ph.D.")
    add_bullet(doc, "Perceived Social Support by Dr. Madhu Asthana & Dr. Kiran Bala Verma")
    add_bullet(doc, "Perceived Stress Scale by Arun K. Singh & Ashish Kumar Singh")

    # Variables
    add_heading(doc, "Variables and Study Design", 1)
    add_paragraph(doc, "Based on the questionnaire, the following variables were used:")
    add_paragraph(doc, "Covariates:", bold=True)
    for v in [
        "Age (IN YEARS)", "Marital Status", "Highest level of education completed", "Current Employment Status", "Number of children", "Type of family structure", "Monthly household income", "Access to healthcare services", "Overall health status", "Are you covered by health insurance?", "Participation in community or social activities"]:
        add_bullet(doc, v)
    add_paragraph(doc, "Independent variables:", bold=True)
    for v in [
        "Functional Impairment", "Withdrawl", "Occupational & Relationship Consequences", "Compulsive Behaviour", "Obsession with Internet", "Internet as a Source of Recreation", "Enhanced Socialization", "Perceived Control of Internet Use", "Internet Usage Total Score"]:
        add_bullet(doc, v)
    add_paragraph(doc, "Dependent variables:", bold=True)
    for v in [
        "Emotional Support", "Informational Support", "Instrumental Support", "Perceived Social Support Total Score", "Power Self Concept - Real Self Concept Score", "Power Self Concept - Ideal Self Concept score", "Power Self Concept – Social Self Concept Score", "Social Self Concept - Real Self Concept Score", "Social Self Concept – Ideal Self Concept Score", "Social Self Concept – Social Self Concept Score", "Ability Self Concept - Real Self Concept Score", "Ability Self Concept - Ideal Self Concept Score", "Ability Self Concept – Social Self Concept Score", "Physical Self Concept - Real Self Concept Score", "Physical Self Concept – Ideal Self Concept Score", "Physical Self Concept – Social Self Concept Score", "Psychological Self Concept - Real Self Concept Score", "Psychological Self Concept – Ideal Self Concept Score", "Psychological Self Concept – Social Self Concept Score", "Real Self Concept Total Score", "Ideal Self Concept Total Score", "Social Self Concept Total Score"]:
        add_bullet(doc, v)

    # Descriptive analysis
    add_heading(doc, "Descriptive Statistical Analysis", 1)
    add_paragraph(doc, "Descriptive statistics were computed for all variables, including means, medians, standard deviations, quartiles, and outlier counts. The following tables summarize the key findings for categorical and numeric variables.")

    # Categorical summary
    add_heading(doc, "Categorical Variables Summary", 2)
    try:
        cat = pd.read_csv('InputData/categorical_summary.csv')
        add_table_from_df(doc, cat.head(10))
        add_paragraph(doc, "Most categorical variables (e.g., Marital Status, Education, Employment) show a reasonable spread across categories, with some (e.g., Number of children) having many missing values. For example, 'Marital Status' is split between Single and Married, and 'Type of family structure' is mostly Nuclear or Joint. Some variables (e.g., health insurance, community participation) show a mix of Yes/No/Not sure responses.")
    except Exception:
        add_paragraph(doc, "(Categorical summary unavailable)")

    # Numeric summary
    add_heading(doc, "Numeric Variables Summary", 2)
    try:
        num = pd.read_csv('InputData/numeric_stats.csv')
        add_table_from_df(doc, num.head(10))
        # Add commentary for each variable
        for _, row in num.iterrows():
            col = row['column']
            mean = row['mean']
            median = row['median']
            std = row['std']
            minv = row['min']
            maxv = row['max']
            n_out = row['n_outliers']
            pct_out = row['pct_outliers']
            skew = (mean - median) / std if std else 0
            comment = f"{col}: mean={mean:.2f}, median={median:.2f}, std={std:.2f}, min={minv}, max={maxv}, outliers={n_out} ({pct_out:.1%}), skewness={skew:.2f}. "
            if pct_out > 0.05:
                comment += "Substantial outliers present. "
            if abs(skew) > 0.5:
                comment += "Distribution is skewed. "
            else:
                comment += "Distribution is approximately symmetric. "
            doc.add_paragraph(comment)
    except Exception:
        add_paragraph(doc, "(Numeric summary unavailable)")

    doc.save('InputData/thesis_descriptive_analysis.docx')
    print('Saved thesis descriptive analysis to InputData/thesis_descriptive_analysis.docx')

if __name__ == '__main__':
    main()
